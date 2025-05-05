import tkinter as tk
from datetime import datetime, date
from tkinter import messagebox, ttk
from tkinter import filedialog
import locale
from typing import Callable # Import Callable for type hinting

# Assume CustomerInfo and RevenueData are imported
from view.db.database import DB_Connector
from view.models import CustomerInfo, RevenueData # Ensure RevenueData is imported

# Import python-docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from view.utils import delete_customer_folder

# --- Define colors --- (Keep consistent)
COLOR_PRIMARY_BLUE = "#3B82F6"
COLOR_ACCENT_GREEN = "#28a745"
COLOR_ACCENT_RED = "#dc3545"
COLOR_ACCENT_TEAL = "#17a2b8"
COLOR_BACKGROUND_LIGHT = "#eef2f7"
COLOR_FRAME_BACKGROUND = "#f8f9fa"
COLOR_MAIN_PANEL_BG = "#ffffff"
COLOR_TEXT_DARK = "#333333"
COLOR_TEXT_MEDIUM = "#555555"
COLOR_BORDER_GRAY = "#cccccc"

# --- Currency Formatting (Keep as is) ---
try:
    locale.setlocale(locale.LC_ALL, 'vi_VN.UTF-8')
except locale.Error:
    try:
        locale.setlocale(locale.LC_ALL, 'Vietnamese_Vietnam.1258')
    except locale.Error:
        print("Warning: Could not set Vietnamese locale for currency formatting.")
        def format_currency(amount):
            return f"{amount:,.0f}".replace(",", ".") + " VND"
    else:
        def format_currency(amount):
            return locale.currency(amount, grouping=True, symbol=" VND")
else:
    def format_currency(amount):
        return locale.currency(amount, grouping=True, symbol=" VND")

if 'format_currency' not in locals():
     def format_currency(amount):
            return f"{amount:,.0f}".replace(",", ".") + " VND"
# --- End Currency Formatting ---


class Checkout(tk.Frame):
    # controller is the App instance
    def __init__(self, parent, controller, revenue_callback: Callable[[RevenueData], None], db_conn: DB_Connector):
        super().__init__(parent, bg=COLOR_BACKGROUND_LIGHT)
        # Store the App instance as the controller
        self.controller = controller
        # Store the callback to send revenue data back to the App
        self.revenue_callback = revenue_callback
        self.db_conn = db_conn

        self.price_per_day = {"VIP": 400000, "Thường": 250000} # Use "Thường" as in customer tab

        self.additional_price_per_day = {
            "Tiền điện": 5000,
            "Tiền nước": 3000,
            "Internet": 10000,
            "Tiền rác": 1000,
            "Vệ sinh": 2000,
        }

        # Panels
        self.leftPanel = tk.Frame(self, bg=COLOR_BACKGROUND_LIGHT)
        self.rightPanel = tk.Frame(self, bg=COLOR_BACKGROUND_LIGHT)
        self.mainPanel = tk.Frame(self, bg=COLOR_MAIN_PANEL_BG, padx=30, pady=30, relief=tk.RAISED, borderwidth=1, highlightbackground=COLOR_BORDER_GRAY, highlightthickness=1)

        # Configure grid weights
        self.columnconfigure(0, weight=1)
        self.columnconfigure(1, weight=5)
        self.columnconfigure(2, weight=1)
        self.rowconfigure(0, weight=1)

        # Layout
        self.leftPanel.grid(row=0, column=0, sticky="nsew")
        self.mainPanel.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        self.rightPanel.grid(row=0, column=2, sticky="nsew")

        # --- Content inside mainPanel ---

        # Title
        self.invoice_label = tk.Label(
            self.mainPanel, text="HÓA ĐƠN THANH TOÁN", font=("Arial", 30, "bold"), bg=COLOR_MAIN_PANEL_BG, fg=COLOR_TEXT_DARK
        )
        self.invoice_label.pack(pady=(0, 20))

        # Customer Info Frame
        self.info_frame = tk.LabelFrame(self.mainPanel, text="THÔNG TIN KHÁCH HÀNG", font=("Arial", 14, "bold"), bg=COLOR_FRAME_BACKGROUND, padx=20, pady=15, bd=1, relief=tk.GROOVE, fg=COLOR_TEXT_MEDIUM)
        self.info_frame.pack(fill="x", padx=10, pady=10)

        info_grid_frame = tk.Frame(self.info_frame, bg=COLOR_FRAME_BACKGROUND)
        info_grid_frame.pack(fill="both", expand=True)

        self.labels = {} # Labels for customer info and room type (read-only)

        # --- Store info fields as a class attribute ---
        self.info_fields = [
            ("Họ tên", "name"), ("Giới tính", "sex"), ("Ngày Sinh", "birthday"),
            ("Quốc tịch", "national"), ("Quê quán", "country"),
        ]
        # --- End NEW ---

        # Use self.info_fields to create labels
        for i, (title, field) in enumerate(self.info_fields):
            tk.Label(info_grid_frame, text=title + ":", font=("Arial", 12, "bold"), bg=COLOR_FRAME_BACKGROUND, fg=COLOR_TEXT_DARK).grid(row=i, column=0, sticky="w", padx=5, pady=3)
            self.labels[field] = tk.Label(info_grid_frame, text="Đang tải...", font=("Arial", 12), bg=COLOR_FRAME_BACKGROUND, fg=COLOR_TEXT_DARK)
            self.labels[field].grid(row=i, column=1, sticky="w", padx=5, pady=3)

        # Room Details Frame
        self.room_frame = tk.LabelFrame(self.mainPanel, text="CHI TIẾT THUÊ PHÒNG", font=("Arial", 14, "bold"), bg=COLOR_FRAME_BACKGROUND, padx=20, pady=15, bd=1, relief=tk.GROOVE, fg=COLOR_TEXT_MEDIUM)
        self.room_frame.pack(fill="x", padx=10, pady=10)

        room_grid_frame = tk.Frame(self.room_frame, bg=COLOR_FRAME_BACKGROUND)
        room_grid_frame.pack(fill="both", expand=True)

        tk.Label(room_grid_frame, text="Loại phòng:", font=("Arial", 12, "bold"), bg=COLOR_FRAME_BACKGROUND, fg=COLOR_TEXT_DARK).grid(row=0, column=0, sticky="w", padx=5, pady=3)
        self.labels["room_type"] = tk.Label(room_grid_frame, text="Đang tải...", font=("Arial", 12), bg=COLOR_FRAME_BACKGROUND, fg=COLOR_TEXT_DARK) # room_type label is also stored here
        self.labels["room_type"].grid(row=0, column=1, sticky="w", padx=5, pady=3)

        tk.Label(room_grid_frame, text="Ngày nhận phòng:", font=("Arial", 12, "bold"), bg=COLOR_FRAME_BACKGROUND, fg=COLOR_TEXT_DARK).grid(row=1, column=0, sticky="w", padx=5, pady=3)
        self.checkin_date_label = tk.Label(room_grid_frame, text="Đang tải...", bg=COLOR_FRAME_BACKGROUND, font=("Arial", 12), fg=COLOR_TEXT_DARK)
        self.checkin_date_label.grid(row=1, column=1, sticky="w", padx=5, pady=3)

        tk.Label(room_grid_frame, text="Ngày trả phòng:", font=("Arial", 12, "bold"), bg=COLOR_FRAME_BACKGROUND, fg=COLOR_TEXT_DARK).grid(row=2, column=0, sticky="w", padx=5, pady=3)
        self.checkout_date_label = tk.Label(room_grid_frame, text="Đang tải...", bg=COLOR_FRAME_BACKGROUND, font=("Arial", 12), fg=COLOR_TEXT_DARK)
        self.checkout_date_label.grid(row=2, column=1, sticky="w", padx=5, pady=3)


        # Separator
        separator_room_to_additional = tk.Frame(self.mainPanel, height=1, bg=COLOR_BORDER_GRAY)
        separator_room_to_additional.pack(fill="x", padx=10, pady=15)

        # Additional Costs Frame (Displaying calculated values)
        self.additional_frame = tk.LabelFrame(self.mainPanel, text="CHI PHÍ PHỤ (Tính theo ngày thuê)", font=("Arial", 14, "bold"), bg=COLOR_FRAME_BACKGROUND, padx=20, pady=15, bd=1, relief=tk.GROOVE, fg=COLOR_TEXT_MEDIUM)
        self.additional_frame.pack(fill="x", padx=10, pady=10)

        additional_grid_frame = tk.Frame(self.additional_frame, bg=COLOR_FRAME_BACKGROUND)
        additional_grid_frame.pack(fill="both", expand=True)

        self.additional_cost_labels = {}
        for i, (cost_name, daily_rate) in enumerate(self.additional_price_per_day.items()):
             tk.Label(additional_grid_frame, text=cost_name + ":", font=("Arial", 12, "bold"), bg=COLOR_FRAME_BACKGROUND, fg=COLOR_TEXT_DARK).grid(row=i, column=0, sticky="w", padx=5, pady=3)
             self.additional_cost_labels[cost_name] = tk.Label(additional_grid_frame, text="Đang tính...", font=("Arial", 12), bg=COLOR_FRAME_BACKGROUND, fg=COLOR_TEXT_DARK)
             self.additional_cost_labels[cost_name].grid(row=i, column=1, sticky="ew", padx=5, pady=3)


        additional_grid_frame.columnconfigure(1, weight=1)


        # Separator
        separator_additional_to_total = tk.Frame(self.mainPanel, height=1, bg=COLOR_BORDER_GRAY)
        separator_additional_to_total.pack(fill="x", padx=10, pady=15)

        # Result Label (Total Cost)
        self.result_label = tk.Label(self.mainPanel, text="Tổng tiền: Đang tính...", font=("Arial", 20, "bold"), fg=COLOR_PRIMARY_BLUE, bg=COLOR_MAIN_PANEL_BG)
        self.result_label.pack(pady=(10, 20))

        # Button Frame
        self.button_frame = tk.Frame(self.mainPanel, bg=COLOR_MAIN_PANEL_BG)
        self.button_frame.pack(pady=(0, 10))

        button_font = ("Arial", 12, "bold")
        button_pady = 8
        button_padx = 15

        # Calculate Button
        self.calculate_btn = tk.Button(
            self.button_frame, text="TÍNH TIỀN", font=button_font, fg="white", bg=COLOR_ACCENT_GREEN,
            activebackground="#1e7e34", activeforeground="white",
            relief=tk.RAISED, padx=button_padx, pady=button_pady, cursor="hand2",
            command=self.load_and_calculate # This will now also be called by refresh_display
        )
        self.calculate_btn.pack(side=tk.LEFT, padx=5)

        # Export Button
        self.export_btn = tk.Button(
            self.button_frame, text="XUẤT FILE WORD", font=button_font, fg="white", bg=COLOR_ACCENT_TEAL,
            activebackground="#117a8b", activeforeground="white",
            relief=tk.RAISED, padx=button_padx, pady=button_pady, cursor="hand2",
            command=self.export_invoice_to_word
        )
        self.export_btn.pack(side=tk.LEFT, padx=5)

        # Finalize Checkout Button (New)
        self.finalize_checkout_btn = tk.Button(
            self.button_frame, text="HOÀN TẤT THANH TOÁN", font=button_font, fg="white", bg=COLOR_PRIMARY_BLUE,
            activebackground="#0056b3", activeforeground="white",
            relief=tk.RAISED, padx=button_padx, pady=button_pady, cursor="hand2",
            command=self.finalize_checkout # New method to finalize and send data
        )
        self.finalize_checkout_btn.pack(side=tk.LEFT, padx=5)


    # --- Helper to process raw date input (string or datetime) into datetime object ---
    def _process_date_input_to_datetime(self, raw_date_value):
         if isinstance(raw_date_value, datetime):
             return raw_date_value # Already a datetime object
         elif isinstance(raw_date_value, date): # Handle date objects too
             return datetime.combine(raw_date_value, datetime.min.time()) # Convert date to datetime
         elif isinstance(raw_date_value, str) and raw_date_value:
             try:
                  # Try parsing from YYYY-MM-DD string
                  return datetime.strptime(raw_date_value, "%Y-%m-%d")
             except ValueError:
                  print(f"Warning: Could not parse date string '{raw_date_value}'. Expected YYYY-MM-DD.")
                  return None
         else:
             return None # None, empty string, or other types are treated as no date


    # --- Calculate Base Room Cost from Duration ---
    def calculate_room_cost_from_duration(self, duration, room_type):
        if duration is None or duration <= 0:
            return 0

        room_rate = self.price_per_day.get(room_type, self.price_per_day.get("Thường", 250000)) # Default to "Thường" if type not found

        return duration * room_rate

    # --- refresh_display method (Called by App when tab is shown) ---
    def refresh_display(self):
        print("[*] Checkout tab refreshing display and recalculating...")
        self.load_and_calculate() # Call the existing method to load and calculate


    # --- load_and_calculate method (Refined Label Updating) ---
    def load_and_calculate(self):
        customer_data = getattr(self.controller, 'current_customer_for_checkout', None)

        if not customer_data:
             # Check if the main window still exists before showing messagebox
             if self.winfo_exists():
                  messagebox.showwarning("Cảnh báo", "Không có dữ liệu khách hàng để tạo hóa đơn.")
             # Clear customer/room info labels (use self.labels keys or self.info_fields)
             for field in self.labels.keys(): # Iterate through existing label keys
                 # Exclude room_type as it's handled separately below
                 if field != "room_type":
                      self.labels[field].config(text="Không có dữ liệu")

             self.checkin_date_label.config(text="N/A")
             self.checkout_date_label.config(text="N/A")
             self.labels["room_type"].config(text="Không có dữ liệu") # Clear room type
             for _, label_widget in self.additional_cost_labels.items():
                  label_widget.config(text="N/A")
             self.result_label.config(text="Tổng tiền: N/A", fg=COLOR_ACCENT_RED)
             return

        # Get raw check-in date data and room type from the customer_data object
        raw_checkin_date = getattr(customer_data, "checkin_date", None)
        room_type = getattr(customer_data, "room_type", "Thường") # Default to "Thường"

        # --- Process Raw Check-in Date ---
        checkin_dt = self._process_date_input_to_datetime(raw_checkin_date)
        duration = 0
        checkin_display = "N/A" # Initialize display string
        checkout_display = datetime.today().strftime("%Y-%m-%d") # Use YYYY-MM-DD format for display

        if checkin_dt:
            checkout_dt = datetime.today() # datetime object for calculation
            if checkout_dt >= checkin_dt:
                # Calculate duration in days, ensure minimum 1 day if check-in is today
                duration = (checkout_dt - checkin_dt).days
                if duration == 0:
                    duration = 1 # Assume at least 1 day if check-in is today
                elif duration < 0:
                    duration = 0 # Should not happen with the check above, but for safety

            else:
                 # Check if the main window still exists before showing messagebox
                 if self.winfo_exists():
                      messagebox.showwarning("Lỗi", "Ngày trả phòng không hợp lệ (trước ngày nhận phòng).")
                 duration = 0 # Reset duration on invalid date range

            checkin_display = checkin_dt.strftime("%Y-%m-%d") # YYYY-MM-DD format for display
        else:
             # Check if the main window still exists before showing messagebox
             if self.winfo_exists():
                  messagebox.showwarning("Cảnh báo", "Ngày nhận phòng không hợp lệ hoặc không có dữ liệu.")
             # checkin_display remains "N/A"


        # --- Update GUI display labels ---

        # Update customer info labels using self.info_fields
        for title, field in self.info_fields:
             value = getattr(customer_data, field, "Không có dữ liệu")
             display_value = str(value) if value is not None else "Không có dữ liệu"

             # Special handling for birthday date format
             if field == "birthday":
                  birthday_dt = self._process_date_input_to_datetime(value)
                  display_value = birthday_dt.strftime("%Y-%m-%d") if birthday_dt else "N/A"

             if field in self.labels: # Ensure label exists before config
                  self.labels[field].config(text=display_value)

        # Update room type label (handled separately as it's not in self.info_fields)
        if "room_type" in self.labels:
             self.labels["room_type"].config(text=str(room_type) if room_type is not None else "Không có dữ liệu")

        # Update check-in and check-out date labels
        self.checkin_date_label.config(text=checkin_display)
        self.checkout_date_label.config(text=checkout_display)


        # --- Calculate Base Room Cost --- (Using duration calculated above)
        base_room_cost = self.calculate_room_cost_from_duration(duration, room_type)


        # --- Calculate and display additional costs --- (Using duration calculated above)
        additional_cost_sum = 0
        for cost_name, daily_rate in self.additional_price_per_day.items():
            cost = daily_rate * duration
            additional_cost_sum += cost
            if cost_name in self.additional_cost_labels:
                self.additional_cost_labels[cost_name].config(text=format_currency(cost))
            # else: (Warning already printed if label not found during setup)


        # Calculate final total
        final_total = base_room_cost + additional_cost_sum

        # Format the final total for display
        formatted_total = format_currency(final_total)

        self.result_label.config(text=f"TỔNG TIỀN: {formatted_total}", fg=COLOR_ACCENT_GREEN if final_total >= 0 else COLOR_ACCENT_RED)

        print("[✓] Calculation complete and display updated.")


    # --- finalize_checkout method (New method to create RevenueData and send back) ---
    def finalize_checkout(self):
        customer_data = getattr(self.controller, 'current_customer_for_checkout', None)

        if not customer_data:
             if self.winfo_exists():
                  messagebox.showwarning("Thiếu thông tin", "Không có dữ liệu khách hàng để hoàn tất thanh toán.")
             return

        try:
             # Re-calculate total to be safe
             raw_checkin_date = getattr(customer_data, "checkin_date", None)
             room_type = getattr(customer_data, "room_type", "Thường")
             checkin_dt = self._process_date_input_to_datetime(raw_checkin_date)
             duration = 0
             if checkin_dt:
                  checkout_dt = datetime.today()
                  if checkout_dt >= checkin_dt:
                       duration = (checkout_dt - checkin_dt).days
                       if duration == 0: duration = 1
                       elif duration < 0: duration = 0
                  else:
                       duration = 0 # Invalid date range

             base_room_cost = self.calculate_room_cost_from_duration(duration, room_type)
             additional_cost_sum = sum(rate * duration for rate in self.additional_price_per_day.values())
             final_total = base_room_cost + additional_cost_sum

        except Exception as e:
             print(f"[!] Error during recalculation for finalize_checkout: {e}")
             if self.winfo_exists():
                  messagebox.showerror("Lỗi Tính toán", f"Không thể tính toán lại tổng tiền để hoàn tất thanh toán:\n{e}")
             return


        # Get the checkout date string (today's date)
        checkout_date_str = datetime.today().strftime("%Y-%m-%d") # YYYY-MM-DD string

        # Create a RevenueData object, passing all required attributes
        revenue_record = RevenueData(
            id=getattr(customer_data, 'id', None),
            name=getattr(customer_data, 'name', None),
            sex=getattr(customer_data, 'sex', None),
            birthday=getattr(customer_data, 'birthday', None), # Should be YYYY-MM-DD string or None
            national=getattr(customer_data, 'national', None),
            country=getattr(customer_data, 'country', None),
            checkin_date=getattr(customer_data, 'checkin_date', None), # Should be YYYY-MM-DD string or None
            room_type=getattr(customer_data, 'room_type', None),
            room_number=getattr(customer_data, 'room_number', None),
            checkout_date=checkout_date_str, # Use the calculated checkout date string
            total_price=final_total # Use the calculated total price
        )

        # Call the callback function provided by the App to add the revenue record
        if self.revenue_callback:
            try:
                self.revenue_callback(revenue_record)
                if self.winfo_exists():
                     messagebox.showinfo("Thành công", "Thanh toán đã hoàn tất và doanh thu đã được ghi nhận.")

                print("[*] Assuming App's revenue_callback handles clearing current customer data.")

                # save revenue to database
                self.db_conn.setRevenueToDatabase(revenue_record)

                # delete customer id folder
                delete_customer_folder(str(revenue_record.id))

                # Clear the display in the Checkout tab
                self.clear_display()

            except Exception as e:
                print(f"[!] Error calling revenue_callback in finalize_checkout: {e}")
                if self.winfo_exists():
                     messagebox.showerror("Lỗi Xử lý Doanh thu", f"Đã xảy ra lỗi khi ghi nhận doanh thu:\n{e}")
        else:
            print("[!] Revenue callback is not set in Checkout tab.")
            if self.winfo_exists():
                 messagebox.showerror("Lỗi Cấu hình", "Không tìm thấy hàm callback để ghi nhận doanh thu.")


    def clear_display(self):
        # Clear customer info labels using self.info_fields
        for title, field in self.info_fields:
             if field in self.labels:
                  self.labels[field].config(text="Đang tải...") # Reset to initial text

        # Clear room type label
        if "room_type" in self.labels:
             self.labels["room_type"].config(text="Đang tải...")

        # Clear date labels
        self.checkin_date_label.config(text="Đang tải...")
        self.checkout_date_label.config(text="Đang tải...")

        # Clear additional cost labels
        for cost_name in self.additional_price_per_day.keys():
             if cost_name in self.additional_cost_labels:
                  self.additional_cost_labels[cost_name].config(text="Đang tính...")

        # Clear total result label
        self.result_label.config(text="Tổng tiền: Đang tính...", fg=COLOR_PRIMARY_BLUE)


    # --- export_invoice_to_word method (Uses data from controller and calculated values) ---
    def export_invoice_to_word(self):
        # Get the current customer data from the controller (App instance)
        customer_data = getattr(self.controller, 'current_customer_for_checkout', None)

        if not customer_data:
            if self.winfo_exists():
                 messagebox.showwarning("Cảnh báo", "Không có dữ liệu khách hàng để xuất hóa đơn.")
            return

        # Ensure calculations have been performed to get duration and total
        # It's safer to re-calculate or store these in instance variables
        try:
             raw_checkin_date = getattr(customer_data, "checkin_date", None)
             room_type = getattr(customer_data, "room_type", "Thường")
             checkin_dt = self._process_date_input_to_datetime(raw_checkin_date)
             duration = 0
             checkin_display = "N/A"
             checkout_display = datetime.today().strftime("%Y-%m-%d") # YYYY-MM-DD for document

             if checkin_dt:
                  checkout_dt = datetime.today()
                  if checkout_dt >= checkin_dt:
                       duration = (checkout_dt - checkin_dt).days
                       if duration == 0: duration = 1
                       elif duration < 0: duration = 0
                  checkin_display = checkin_dt.strftime("%Y-%m-%d")

             base_room_cost = self.calculate_room_cost_from_duration(duration, room_type)
             additional_cost_sum = sum(rate * duration for rate in self.additional_price_per_day.values())
             final_total = base_room_cost + additional_cost_sum

        except Exception as e:
             print(f"[!] Error during recalculation for export_invoice_to_word: {e}")
             if self.winfo_exists():
                  messagebox.showerror("Lỗi Tính toán", f"Không thể tính toán lại tổng tiền để xuất hóa đơn:\n{e}")
             return


        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")],
            title="Lưu Hóa Đơn Dạng Word"
        )

        if not filepath:
            return

        try:
            document = Document()
            title = document.add_paragraph("HÓA ĐƠN THANH TOÁN")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.style = 'Heading 1'

            # Add Customer Info
            document.add_paragraph("THÔNG TIN KHÁCH HÀNG", style='Heading 2')
            # Use self.info_fields
            for title, field in self.info_fields:
                value = getattr(customer_data, field, "Không có dữ liệu")
                display_value = str(value) if value is not None else "Không có dữ liệu"

                # Format dates for document (expecting YYYY-MM-DD string or datetime)
                if field in ["birthday"] and isinstance(value, (datetime, date)):
                     birthday_dt = self._process_date_input_to_datetime(value)
                     display_value = birthday_dt.strftime("%Y-%m-%d") if birthday_dt else "N/A"


                p = document.add_paragraph()
                p.add_run(title + ": ").bold = True
                p.add_run(display_value)


            # Add Room Details (using calculated duration/dates)
            document.add_paragraph("CHI TIẾT THUÊ PHÒNG", style='Heading 2')

            p_room = document.add_paragraph()
            p_room.add_run("Loại phòng: ").bold = True
            p_room.add_run(str(room_type) if room_type is not None else "Không có dữ liệu")

            p_checkin = document.add_paragraph()
            p_checkin.add_run("Ngày nhận phòng: ").bold = True
            p_checkin.add_run(checkin_display)

            p_checkout = document.add_paragraph()
            p_checkout.add_run("Ngày trả phòng: ").bold = True
            p_checkout.add_run(checkout_display)

            p_duration = document.add_paragraph()
            p_duration.add_run("Tổng số ngày thuê: ").bold = True
            p_duration.add_run(f"{duration} ngày")


            # Add Additional Costs to Word
            document.add_paragraph("CHI PHÍ PHỤ", style='Heading 2')

            # Use the calculated costs from the load_and_calculate logic
            for cost_name, daily_rate in self.additional_price_per_day.items():
                 cost = daily_rate * duration # Recalculate cost for the document

                 p_cost = document.add_paragraph()
                 p_cost.add_run(cost_name + ": ").bold = True
                 p_cost.add_run(format_currency(cost))


            # Add Total Cost
            document.add_paragraph("")
            p_total = document.add_paragraph()
            p_total.add_run("TỔNG TIỀN: ").bold = True
            p_total.add_run(format_currency(final_total)).bold = True
            p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            document.save(filepath)
            if self.winfo_exists():
                 messagebox.showinfo("Thành công", f"Đã xuất hóa đơn ra file:\n{filepath}")

        except Exception as e:
            print(f"[!] Error during Word export: {e}")
            if self.winfo_exists():
                 messagebox.showerror("Lỗi", f"Đã xảy ra lỗi khi xuất file Word:\n{e}")

